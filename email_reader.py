"""
Модуль для чтения писем в разных форматах и извлечения текста
"""
import email
from email import policy
from email.message import EmailMessage
from pathlib import Path
import re
import base64
import tempfile
import os
from bs4 import BeautifulSoup
import magic  # python-magic-bin

# Опциональные импорты (могут не быть установлены)
try:
    from PIL import Image
    import pytesseract
    # Укажите путь к Tesseract если он не в PATH
    # pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠ OCR не доступен. Текст из изображений извлекаться не будет.")
    print("  Установите: pip install pytesseract Pillow")
    print("  И скачайте Tesseract: https://github.com/UB-Mannheim/tesseract/wiki")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class EmailReader:
    """Читает письма в разных форматах и извлекает весь текст"""
    
    def __init__(self):
        self.supported_formats = ['.eml', '.txt', '.html', '.htm']
    
    def read_file(self, filepath: Path) -> str:
        """
        Основной метод: читает файл письма и возвращает весь текст
        """
        suffix = filepath.suffix.lower()
        
        if suffix == '.eml':
            return self._read_eml(filepath)
        elif suffix in ['.txt', '.html', '.htm']:
            return self._read_text_file(filepath)
        else:
            # Пробуем прочитать как текст
            return self._read_text_file(filepath)
    
    def _read_eml(self, filepath: Path) -> str:
        """Читает .eml файл и извлекает текст из всех частей"""
        with open(filepath, 'rb') as f:
            msg = email.message_from_binary_file(f, policy=policy.default)
        
        # Собираем заголовки
        text_parts = []
        subject = msg.get('subject', '')
        sender = msg.get('from', '')
        
        if subject:
            text_parts.append(f"Subject: {subject}")
        if sender:
            text_parts.append(f"From: {sender}")
        
        # Извлекаем текст из тела письма
        body_text = self._extract_body(msg)
        if body_text:
            text_parts.append(body_text)
        
        # Извлекаем текст из вложений
        attachment_texts = self._extract_attachments(msg)
        if attachment_texts:
            text_parts.append("--- Attachments ---")
            text_parts.extend(attachment_texts)
        
        return "\n".join(text_parts)
    
    def _extract_body(self, msg) -> str:
        """Извлекает текст из тела письма (поддерживает multipart, HTML, plain text)"""
        text_parts = []
        
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                disposition = str(part.get('Content-Disposition', ''))
                
                # Пропускаем вложения (их обработаем отдельно)
                if 'attachment' in disposition:
                    continue
                
                # Извлекаем текст
                text = self._get_part_text(part, content_type)
                if text:
                    text_parts.append(text)
        else:
            content_type = msg.get_content_type()
            text = self._get_part_text(msg, content_type)
            if text:
                text_parts.append(text)
        
        return "\n".join(text_parts)
    
    def _get_part_text(self, part, content_type: str) -> str:
        """Извлекает текст из части письма"""
        try:
            payload = part.get_payload(decode=True)
            if not payload:
                return ""
            
            charset = part.get_content_charset() or 'utf-8'
            
            if content_type == 'text/plain':
                return payload.decode(charset, errors='ignore')
            
            elif content_type == 'text/html':
                html_text = payload.decode(charset, errors='ignore')
                return self._html_to_text(html_text)
            
            elif content_type.startswith('image/'):
                # Извлекаем текст из картинки через OCR
                return self._image_to_text(payload)
            
            else:
                # Пробуем как текст
                return payload.decode(charset, errors='ignore')
                
        except Exception as e:
            return f"[Error reading part: {e}]"
    
    def _extract_attachments(self, msg) -> list:
        """Извлекает текст из вложений"""
        texts = []
        
        if not msg.is_multipart():
            return texts
        
        for part in msg.walk():
            disposition = str(part.get('Content-Disposition', ''))
            if 'attachment' not in disposition:
                continue
            
            filename = part.get_filename() or 'unknown'
            content_type = part.get_content_type()
            
            try:
                payload = part.get_payload(decode=True)
                if not payload:
                    continue
                
                text = self._process_attachment(payload, filename, content_type)
                if text:
                    texts.append(f"[Attachment: {filename}]\n{text}")
                    
            except Exception as e:
                texts.append(f"[Attachment: {filename} - Error: {e}]")
        
        return texts
    
    def _process_attachment(self, payload: bytes, filename: str, content_type: str) -> str:
        """Обрабатывает вложение в зависимости от типа"""
        
        # Сохраняем во временный файл
        suffix = Path(filename).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(payload)
            tmp_path = tmp.name
        
        try:
            # Обработка по типу файла
            if content_type == 'text/plain' or suffix in ['.txt', '.csv']:
                return payload.decode('utf-8', errors='ignore')
            
            elif content_type == 'text/html' or suffix in ['.html', '.htm']:
                return self._html_to_text(payload.decode('utf-8', errors='ignore'))
            
            elif content_type.startswith('image/') or suffix in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
                return self._image_to_text(payload)
            
            elif suffix == '.pdf':
                return self._pdf_to_text(tmp_path)
            
            elif suffix in ['.docx', '.doc']:
                return self._docx_to_text(tmp_path)
            
            elif suffix in ['.xlsx', '.xls']:
                return self._excel_to_text(tmp_path)
            
            else:
                # Пробуем прочитать как текст
                return payload.decode('utf-8', errors='ignore')
                
        finally:
            # Удаляем временный файл
            try:
                os.unlink(tmp_path)
            except:
                pass
    
    def _html_to_text(self, html: str) -> str:
        """Извлекает текст из HTML, сохраняя структуру"""
        try:
            soup = BeautifulSoup(html, 'lxml')
            
            # Удаляем скрипты и стили
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Извлекаем текст
            text = soup.get_text(separator='\n', strip=True)
            
            # Убираем множественные переносы строк
            text = re.sub(r'\n\s*\n', '\n', text)
            
            return text
        except:
            # Если lxml не работает, пробуем встроенный парсер
            soup = BeautifulSoup(html, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text(separator='\n', strip=True)
    
    def _image_to_text(self, image_data: bytes) -> str:
        """Извлекает текст из изображения через OCR"""
        if not OCR_AVAILABLE:
            return "[Image - OCR not available]"
        
        try:
            from io import BytesIO
            image = Image.open(BytesIO(image_data))
            text = pytesseract.image_to_string(image, lang='eng+rus+chi_sim')
            return text.strip()
        except Exception as e:
            return f"[Image OCR error: {e}]"
    
    def _pdf_to_text(self, filepath: str) -> str:
        """Извлекает текст из PDF"""
        if not PDF_AVAILABLE:
            return "[PDF - pdfplumber not available]"
        
        try:
            texts = []
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        texts.append(page_text)
                    
                    # Извлекаем текст из таблиц
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            table_text = self._format_table(table)
                            texts.append(table_text)
            
            return "\n".join(texts)
        except Exception as e:
            return f"[PDF error: {e}]"
    
    def _docx_to_text(self, filepath: str) -> str:
        """Извлекает текст из Word документа"""
        if not DOCX_AVAILABLE:
            return "[DOCX - python-docx not available]"
        
        try:
            doc = docx.Document(filepath)
            texts = []
            
            # Текст из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            # Текст из таблиц
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                texts.append(self._format_table(table_data))
            
            return "\n".join(texts)
        except Exception as e:
            return f"[DOCX error: {e}]"
    
    def _excel_to_text(self, filepath: str) -> str:
        """Извлекает текст из Excel файла"""
        if not EXCEL_AVAILABLE:
            return "[Excel - openpyxl not available]"
        
        try:
            wb = openpyxl.load_workbook(filepath, data_only=True)
            texts = []
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                texts.append(f"[Sheet: {sheet_name}]")
                
                table_data = []
                for row in ws.iter_rows(values_only=True):
                    table_data.append([str(cell) if cell is not None else '' for cell in row])
                
                texts.append(self._format_table(table_data))
            
            return "\n".join(texts)
        except Exception as e:
            return f"[Excel error: {e}]"
    
    def _format_table(self, table: list) -> str:
        """Форматирует таблицу в текстовый вид"""
        if not table:
            return ""
        
        lines = []
        for row in table:
            # Соединяем ячейки через разделитель
            row_text = " | ".join(str(cell) for cell in row if cell)
            if row_text.strip():
                lines.append(row_text)
        
        return "\n".join(lines)
    
    def _read_text_file(self, filepath: Path) -> str:
        """Читает простой текстовый файл"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'gb2312', 'gbk', 'windows-1251']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    content = f.read()
                
                # Если это HTML - извлекаем текст
                if filepath.suffix in ['.html', '.htm']:
                    return self._html_to_text(content)
                
                return content
            except:
                continue
        
        # Последняя попытка
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


# Тестирование
if __name__ == "__main__":
    reader = EmailReader()
    
    # Тест с простым текстом
    test_text = reader._html_to_text("<html><body><h1>Hello</h1><p>This is a test</p></body></html>")
    print("HTML test:", test_text)
    
    print("\nEmailReader готов к работе!")